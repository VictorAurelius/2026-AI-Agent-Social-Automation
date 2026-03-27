# modules/book-translation/scripts/tests/test_renderer.py
"""Tests for markdown-to-DOCX renderer module."""
import pytest
from pathlib import Path
from docx import Document
from lib.renderer import render_markdown_to_docx, parse_markdown_elements


class TestParseMarkdownElements:
    def test_heading_level_1(self):
        elements = parse_markdown_elements("# Chapter Title\n")
        assert elements[0]["type"] == "heading"
        assert elements[0]["level"] == 1
        assert elements[0]["text"] == "Chapter Title"

    def test_heading_level_2(self):
        elements = parse_markdown_elements("## Section Title\n")
        assert elements[0]["type"] == "heading"
        assert elements[0]["level"] == 2

    def test_paragraph(self):
        elements = parse_markdown_elements("A regular paragraph.\n")
        assert elements[0]["type"] == "paragraph"
        assert elements[0]["text"] == "A regular paragraph."

    def test_bold_text(self):
        elements = parse_markdown_elements("This has **bold** text.\n")
        assert elements[0]["type"] == "paragraph"
        runs = elements[0]["runs"]
        assert any(r["bold"] for r in runs)

    def test_italic_text(self):
        elements = parse_markdown_elements("This has *italic* text.\n")
        assert elements[0]["type"] == "paragraph"
        runs = elements[0]["runs"]
        assert any(r["italic"] for r in runs)

    def test_page_break(self):
        elements = parse_markdown_elements("---\n")
        assert elements[0]["type"] == "page_break"

    def test_blockquote(self):
        elements = parse_markdown_elements("> A quoted passage.\n")
        assert elements[0]["type"] == "blockquote"
        assert elements[0]["text"] == "A quoted passage."

    def test_mixed_content(self):
        md = "# Title\n\nParagraph one.\n\nParagraph two with **bold**.\n\n---\n\n# Next Chapter\n"
        elements = parse_markdown_elements(md)
        types = [e["type"] for e in elements]
        assert types[0] == "heading"
        assert "paragraph" in types
        assert "page_break" in types


class TestRenderMarkdownToDocx:
    def test_creates_docx_file(self, tmp_path):
        md = "# Chapter 1\n\nHello world.\n"
        output = tmp_path / "test.docx"
        render_markdown_to_docx(md, output)
        assert output.exists()

    def test_docx_has_heading(self, tmp_path):
        md = "# My Chapter\n\nContent here.\n"
        output = tmp_path / "test.docx"
        render_markdown_to_docx(md, output)
        doc = Document(str(output))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 1
        assert "My Chapter" in headings[0].text

    def test_docx_has_paragraphs(self, tmp_path):
        md = "# Title\n\nFirst paragraph.\n\nSecond paragraph.\n"
        output = tmp_path / "test.docx"
        render_markdown_to_docx(md, output)
        doc = Document(str(output))
        normal_paras = [p for p in doc.paragraphs if p.style.name == "Normal" and p.text]
        assert len(normal_paras) >= 2

    def test_docx_has_bold(self, tmp_path):
        md = "# Title\n\nThis is **bold text** here.\n"
        output = tmp_path / "test.docx"
        render_markdown_to_docx(md, output)
        doc = Document(str(output))
        bold_runs = [run for p in doc.paragraphs for run in p.runs if run.bold]
        assert len(bold_runs) >= 1
        assert "bold text" in bold_runs[0].text

    def test_docx_font_settings(self, tmp_path):
        md = "# Title\n\nContent.\n"
        output = tmp_path / "test.docx"
        render_markdown_to_docx(md, output, font_name="Times New Roman", font_size=12)
        doc = Document(str(output))
        for p in doc.paragraphs:
            for run in p.runs:
                if run.font.name:
                    assert run.font.name == "Times New Roman"
